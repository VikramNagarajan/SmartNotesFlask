from flask import Flask, render_template, url_for, flash, redirect, request, send_from_directory
from flask_bootstrap import Bootstrap
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField
from wtforms.validators import InputRequired, Email, Length, EqualTo
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
import os

app = Flask(__name__)
app.config['SECRET_KEY'] = os.urandom(12)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
UPLOAD_FOLDER = 'C:\\owner\\Desktop\\test\\uploadfiles'
ALLOWED_EXTENSIONS = {'pdf'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
Bootstrap(app)
db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(15), unique=True)
    email = db.Column(db.String(50), unique=True)
    password = db.Column(db.String(80))
    notes = db.relationship('Note', backref='creator', lazy=True)

class Note(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    noteFile = db.Column(db.String(150))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))



@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


class LoginForm(FlaskForm):
    username = StringField('Username', validators=[InputRequired(), Length(min=4, max=15)])
    password = PasswordField('Password', validators=[InputRequired(), Length(min=8, max=80)])

class RegisterForm(FlaskForm):
    email = username = StringField('email', validators=[InputRequired(), Email(message='Invalid Email'), Length(max=50)])
    username = StringField('username', validators=[InputRequired(), Length(min=4, max=15)])
    password = PasswordField('password', validators=[InputRequired(), Length(min=8, max=80)])

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user:
            if check_password_hash(user.password, form.password.data):
                login_user(user)
                return redirect(url_for('dashboard'))
            else:
                flash('Either the username or password is incorrect.', category='danger')
        else:
            flash('Either the username or password is incorrect.', category='danger')
        
    return render_template('login.html', form=form)

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = RegisterForm()
    if form.validate_on_submit():
        hashed_password = generate_password_hash(form.password.data, method='sha256')
        try:
            new_user = User(username=form.username.data, email=form.email.data, password=hashed_password)
            db.session.add(new_user)
            db.session.commit()

            return redirect(url_for('index'))
        except:
            flash('Either the username or email has already been taken.', category='danger')
        #return '<h1>' + form.username.data + ' ' + form.email.data + ' ' + form.password.data + '</h1>'
    return render_template('signup.html', form=form)

@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', name=current_user.username)


@app.route('/dashstep', methods=['GET', 'POST'])
@login_required
def dashstep():
    return render_template('dashstep.html')

@app.route('/new', methods=['GET', 'POST'])
@login_required
def new():
    return render_template('new.html')

@app.route('/newpdf', methods=['GET', 'POST'])
@login_required
def newpdf():
    return render_template('newpdf.html')

@app.route('/notes', methods=['GET', 'POST'])
@login_required
def notes():
    from PyDictionary import PyDictionary
    from summa import keywords
    from summa.summarizer import summarize
    import nltk
    from nltk.tokenize import sent_tokenize
    from newspaper import Article
    from docx import Document
    url = str(request.form['link'])
    a = Article(url)
    a.download()
    a.parse()
    f = a.text
    b = a.title
    a = a.text
    a = keywords.keywords(a)
    dictionary = PyDictionary()
    a = a.split('\n')
    a1 = []
    for i in a:
        x = i.split(' ')
        for j in x:
                a1.append(j)
    a1.sort(key=lambda s: len(s))
    a1.reverse()
    try:
        a1 = a1[:20]
    except:
        pass
    a = set(a1)
    a = tuple(a1)
    a1 = []
    for i in range(10):
        try:
            a1.append(a[i])
        except:
            pass
    from nltk.stem import WordNetLemmatizer 
    lemmatizer = WordNetLemmatizer() 
    a = a1
    a1 = []
    for i in a:
        a1.append(lemmatizer.lemmatize(i))
    a = list(set(a1))
    a1 = a
    a = [dictionary.meaning(i) for i in a1]

    z = sent_tokenize(summarize(f, ratio=0.25))

    doc = Document()
    doc.add_heading('Notes for ' + b, 0)
    for i in z:
        doc.add_paragraph(i)
    doc.add_heading('Vocab Words from ' + b, 0)
    for i in range(len(a)):
        c = doc.add_paragraph(str(i+1) + ') ')
        c.add_run(a1[i]).bold = True
        c.add_run(': ')
        d = str(list(a[i].values()))
        d = d.replace('[', '')
        d = d.replace(']', '')
        c.add_run(d)
        g =doc.add_paragraph('')
        g.add_run('Synonyms for ')
        g.add_run(a1[i].upper() + ': ').bold = True
        from datamuse import datamuse
        api = datamuse.Datamuse()
        s = api.words(ml=a1[i], max=10)
        s1 = []
        for i in s:
            for j in i:
                if j == 'word':
                    s1.append(i[j])
        g.add_run(str(s1).replace('[', '').replace(']', '').replace("'", '')).italic = True
    whitelist = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')
    fileName = b.replace(' ', '') 
    fileName = ''.join(filter(whitelist.__contains__, fileName))
    fileName += '.docx'
    doc.save(fileName)
    import cloudmersive_convert_api_client
    from cloudmersive_convert_api_client.rest import ApiException
    configuration = cloudmersive_convert_api_client.Configuration()
    configuration.api_key['Apikey'] = 'f0c513bc-8c00-4491-830e-3e83b015feb6'
    api_instance = cloudmersive_convert_api_client.ConvertDocumentApi(cloudmersive_convert_api_client.ApiClient(configuration))
    try:
        # Convert Word DOCX Document to PDF
        api_response = api_instance.convert_document_docx_to_pdf(fileName)
        file = open('static/' + fileName.replace('.docx', '.pdf'), 'wb')
        file.write(api_response)
        file.close()
    except ApiException as e:
        print("Exception when calling ConvertDocumentApi->convert_document_docx_to_pdf: %s\n" % e)
    myFile = fileName.replace('.docx', '.pdf')
    myFile2 = myFile
    note = Note(noteFile=str(myFile2), creator=current_user)
    db.session.add(note)
    db.session.commit()
    myFile = url_for('.static', filename=myFile)
    return render_template('notes.html', myFile=myFile)


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



@app.route('/notespdf', methods=['POST', 'GET'])
@login_required
def notespdf():
    import os
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(f'static\{filename}')
            s = request.form['start']
            e = request.form['end']
          
            return redirect(url_for('uploaded_file',
                                    filename=filename, s=s, e=e))
    return render_template('notespdf.html')

@app.route('/uploads/<filename><s><e>')
def uploaded_file(filename, s, e):
    import fitz
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
    pdffile = filename
    doc = fitz.open('static' + '/' + filename)
    for i in range(int(s)-1, int(e)):
        page = doc.loadPage(i)  # number of page
        pix = page.getPixmap()
        output = "outfile"+str(i)+".png"
        pix.writePNG(output)
    x = ''
    for i in range(int(s)-1, int(e)):
        x += pytesseract.image_to_string(f'outfile{str(i)}.png')
    from PyDictionary import PyDictionary
    from summa import keywords
    from summa.summarizer import summarize
    import nltk
    from nltk.tokenize import sent_tokenize
    from docx import Document
    f = x
    b = str(filename.replace('.pdf', ''))
    a = x
    a = keywords.keywords(a)
    dictionary = PyDictionary()
    a = a.split('\n')
    a1 = []
    for i in a:
        x = i.split(' ')
        for j in x:
                a1.append(j)
    a1.sort(key=lambda s: len(s))
    a1.reverse()
    try:
        a1 = a1[:20]
    except:
        pass
    a = set(a1)
    a = tuple(a1)
    a1 = []
    for i in range(10):
        try:
            a1.append(a[i])
        except:
            pass
    from nltk.stem import WordNetLemmatizer 
    lemmatizer = WordNetLemmatizer() 
    a = a1
    a1 = []
    for i in a:
        a1.append(lemmatizer.lemmatize(i))
    a = list(set(a1))
    a1 = a
    a = [dictionary.meaning(i) for i in a1]

    z = sent_tokenize(summarize(f, ratio=0.25))

    doc = Document()
    doc.add_heading('Notes for ' + b, 0)
    for i in z:
        doc.add_paragraph(i)
    doc.add_heading('Vocab Words from ' + b, 0)
    for i in range(len(a)):
        c = doc.add_paragraph(str(i+1) + ') ')
        c.add_run(a1[i]).bold = True
        c.add_run(': ')
        d = str(list(a[i].values()))
        d = d.replace('[', '')
        d = d.replace(']', '')
        c.add_run(d)
        g =doc.add_paragraph('')
        g.add_run('Synonyms for ')
        g.add_run(a1[i].upper() + ': ').bold = True
        from datamuse import datamuse
        api = datamuse.Datamuse()
        s = api.words(ml=a1[i], max=10)
        s1 = []
        for i in s:
            for j in i:
                if j == 'word':
                    s1.append(i[j])
        g.add_run(str(s1).replace('[', '').replace(']', '').replace("'", '')).italic = True
    whitelist = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')
    fileName = b.replace(' ', '') 
    fileName = ''.join(filter(whitelist.__contains__, fileName))
    fileName += '.docx'
    doc.save(fileName)
    import cloudmersive_convert_api_client
    from cloudmersive_convert_api_client.rest import ApiException
    configuration = cloudmersive_convert_api_client.Configuration()
    configuration.api_key['Apikey'] = 'f0c513bc-8c00-4491-830e-3e83b015feb6'
    api_instance = cloudmersive_convert_api_client.ConvertDocumentApi(cloudmersive_convert_api_client.ApiClient(configuration))
    try:
        # Convert Word DOCX Document to PDF
        api_response = api_instance.convert_document_docx_to_pdf(fileName)
        file = open('static/' + fileName.replace('.docx', '.pdf'), 'wb')
        file.write(api_response)
        file.close()
    except ApiException as e:
        print("Exception when calling ConvertDocumentApi->convert_document_docx_to_pdf: %s\n" % e)
    myFile = fileName.replace('.docx', '.pdf')
    myFile2 = myFile
    note = Note(noteFile=str(myFile2), creator=current_user)
    db.session.add(note)
    db.session.commit()
    myFile = url_for('.static', filename=myFile)
    return render_template('notes.html', myFile=myFile)


@app.route('/old')
@login_required
def old():
    notes = Note.query.filter_by(user_id=current_user.id).all()
    return render_template('old.html', notes=notes)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
