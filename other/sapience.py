from PyDictionary import PyDictionary
from summa import keywords
from summa.summarizer import summarize
import nltk
from nltk.tokenize import sent_tokenize
from newspaper import Article
from docx import Document
import os
url = input('hi ')
a = Article(url)
a.download()
a.parse()
f = a.text
b = a.title
a = a.text
a = keywords.keywords(a)
dictionary = PyDictionary()
a = a.split('\n')
a1 = []
for i in a:
  x = i.split(' ')
  for j in x:
        a1.append(j)
a1.sort(key=lambda s: len(s))
a1.reverse()
try:
    a1 = a1[:20]
except:
    pass
a = set(a1)
a = tuple(a1)
a1 = []
for i in range(10):
  try:
    a1.append(a[i])
  except:
    pass
from nltk.stem import WordNetLemmatizer 
lemmatizer = WordNetLemmatizer() 
a = a1
a1 = []
for i in a:
  a1.append(lemmatizer.lemmatize(i))
a = list(set(a1))
a1 = a
a = [dictionary.meaning(i) for i in a1]

z = sent_tokenize(summarize(f, ratio=0.1))

doc = Document()
doc.add_heading('Notes for ' + b, 0)
for i in z:
    doc.add_paragraph(i)
doc.add_heading('Vocab Words from ' + b, 0)
for i in range(len(a)):
    c = doc.add_paragraph(str(i+1) + ') ')
    c.add_run(a1[i]).bold = True
    c.add_run(': ')
    d = str(list(a[i].values()))
    d = d.replace('[', '')
    d = d.replace(']', '')
    c.add_run(d)
    g =doc.add_paragraph('')
    g.add_run('Synonyms for ')
    g.add_run(a1[i].upper() + ': ').bold = True
    from datamuse import datamuse
    api = datamuse.Datamuse()
    s = api.words(ml=a1[i], max=10)
    s1 = []
    for i in s:
      for j in i:
        if j == 'word':
          s1.append(i[j])
    g.add_run(str(s1).replace('[', '').replace(']', '').replace("'", '')).italic = True
whitelist = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')
fileName = b.replace(' ', '') 
fileName = ''.join(filter(whitelist.__contains__, fileName))
fileName += '.docx'
doc.save(fileName)
import cloudmersive_convert_api_client
from cloudmersive_convert_api_client.rest import ApiException

configuration = cloudmersive_convert_api_client.Configuration()
configuration.api_key['Apikey'] = 'f0c513bc-8c00-4491-830e-3e83b015feb6'


# create an instance of the API class
api_instance = cloudmersive_convert_api_client.ConvertDocumentApi(cloudmersive_convert_api_client.ApiClient(configuration))
input_file = '/path/to/file' # file | Input file to perform the operation on.

try:
    # Convert Word DOCX Document to PDF
    api_response = api_instance.convert_document_docx_to_pdf(fileName)
    file = open(fileName.replace('.docx', '.pdf'), 'wb')
    file.write(api_response)
    file.close()
except ApiException as e:
    print("Exception when calling ConvertDocumentApi->convert_document_docx_to_pdf: %s\n" % e)
import os
os.startfile(fileName.replace('.docx', '.pdf'))